import os
import io
import shutil
import uuid
import urllib.parse
from typing import Optional, List
from datetime import datetime, timedelta

from fastapi import (
    FastAPI, Request, Form, UploadFile, File, Depends, HTTPException, Query
)
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.middleware.sessions import SessionMiddleware
from sqlalchemy.orm import Session
from sqlalchemy import func

from database import SessionLocal
import models
import crud
from auth import hash_password, verify_password, get_db, get_current_user, require_admin
import uvicorn

from email_utils import send_booking_confirmation, send_booking_cancellation

from review_utils import can_user_review_quest, get_user_bookings_for_review, create_review, get_quest_reviews, get_quest_average_rating

# Импорты для генерации отчетов Word
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# --- Функция для получения текущего времени ---
def now_with_tz():
    return datetime.now().astimezone()


def naive_now():
    return datetime.now().replace(tzinfo=None)


# --- Подготовка директорий ---
os.makedirs("static/uploads", exist_ok=True)
os.makedirs("static/images", exist_ok=True)
os.makedirs("templates", exist_ok=True)

app = FastAPI()
app.add_middleware(SessionMiddleware, secret_key="!secret_dev_change_me!")

# --- Статика и шаблоны ---
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

print("✅ Подключение к существующей базе данных")


# --- Хелперы ---
def save_upload(file: UploadFile) -> str:
    ext = os.path.splitext(file.filename)[1]
    safe_name = f"{uuid.uuid4().hex}{ext}"
    dest = os.path.join("static", "uploads", safe_name)
    with open(dest, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    return f"uploads/{safe_name}"


# --- Маршруты ---
@app.get("/", response_class=HTMLResponse)
def index(
        request: Request,
        q: Optional[str] = None,
        genre: Optional[List[str]] = Query(None),  # <-- Важно: Query(None)
        difficulty: Optional[List[str]] = Query(None),  # <-- Важно: Query(None)
        fear_level: Optional[str] = None,
        players: Optional[str] = None,
        sort: Optional[str] = None,
        skip: int = 0,
        db: Session = Depends(get_db)
):
    # Добавьте отладочный вывод
    print(f"DEBUG - genre: {genre}")
    print(f"DEBUG - difficulty: {difficulty}")

    try:
        filters = {
            "q": q,
            "genre": genre,
            "difficulty": difficulty,
            "fear_level": fear_level,
            "players": players,
            "sort": sort
        }
        quests = crud.get_quests(db, skip=skip, limit=15, filters=filters)

        # Получаем общее количество активных квестов
        total_quests = db.query(models.Quest).filter(models.Quest.is_active == True).count()

        user = None
        user_bookings_count = 0

        try:
            user = get_current_user(request, db)
            if user:
                user_bookings_count = db.query(models.Booking).filter(
                    models.Booking.user_id == user.id
                ).count()
        except:
            pass

    except Exception as e:
        print(f"Ошибка при загрузке квестов: {e}")
        quests = []
        total_quests = 0
        user = None
        user_bookings_count = 0

    return templates.TemplateResponse("index.html", {
        "request": request,
        "quests": quests,
        "user": user,
        "skip": skip,
        "now": naive_now,
        "total_quests": total_quests,
        "user_bookings_count": user_bookings_count
    })


@app.get("/quest/{quest_id}", response_class=HTMLResponse)
def quest_detail(request: Request, quest_id: int, db: Session = Depends(get_db)):
    quest = crud.get_quest(db, quest_id)
    if not quest:
        raise HTTPException(status_code=404, detail="Quest not found")

    try:
        booked_slots = crud.get_booked_slots_for_date(db, quest_id, naive_now().strftime('%Y-%m-%d'))
    except:
        booked_slots = []

    try:
        user = get_current_user(request, db)
    except:
        user = None

    return templates.TemplateResponse("quest_detail.html", {
        "request": request,
        "quest": quest,
        "user": user,
        "booked_slots": booked_slots,
        "now": naive_now
    })


@app.get("/api/available-slots")
def get_available_slots(quest_id: int, date: str, db: Session = Depends(get_db)):
    try:
        booked_slots = crud.get_booked_slots_for_date(db, quest_id, date)
        return JSONResponse(booked_slots)
    except Exception as e:
        print(f"Error getting available slots: {e}")
        return JSONResponse([], status_code=500)


@app.post("/book")
def book(request: Request, quest_id: int = Form(...), date: str = Form(...),
         timeslot: str = Form(...), payment_method: str = Form(...),
         db: Session = Depends(get_db)):
    try:
        user = get_current_user(request, db)

        # Проверяем, что выбран способ оплаты
        if not payment_method:
            return JSONResponse({"success": False, "message": "Выберите способ оплаты"}, status_code=400)

        # Проверяем допустимые способы оплаты
        if payment_method not in ['card', 'sbp']:
            return JSONResponse({"success": False, "message": "Недопустимый способ оплаты"}, status_code=400)

        # Проверяем, что бронирование не менее чем за 24 часа
        booking_datetime = datetime.strptime(f"{date} {timeslot}", "%Y-%m-%d %H:%M")
        current_time = naive_now()
        hours_diff = (booking_datetime - current_time).total_seconds() / 3600

        if hours_diff < 24:
            return JSONResponse({
                "success": False,
                "message": f"Бронирование возможно не менее чем за 24 часа. До начала квеста осталось {int(hours_diff)} часов"
            }, status_code=400)

        booking = crud.create_booking(db, user_id=user.id, quest_id=quest_id,
                                      date=date, timeslot=timeslot,
                                      payment_method=payment_method)
        if not booking:
            return JSONResponse({"success": False, "message": "Слот уже занят или произошла ошибка"}, status_code=400)

        # Возвращаем информацию о предоплате
        return JSONResponse({
            "success": True,
            "message": f"Бронь создана! Требуется внести предоплату {booking.prepayment}₽",
            "booking_id": booking.id,
            "prepayment": booking.prepayment,
            "total_price": booking.total_price,
            "payment_method": payment_method
        })
    except HTTPException:
        return JSONResponse({"success": False, "message": "Необходимо авторизоваться"}, status_code=401)
    except Exception as e:
        print(f"Booking error: {e}")
        return JSONResponse({"success": False, "message": f"Ошибка: {str(e)}"}, status_code=500)


@app.get("/my-bookings", response_class=HTMLResponse)
def my_bookings(request: Request, db: Session = Depends(get_db)):
    try:
        user = get_current_user(request, db)
        bookings = crud.get_user_bookings(db, user.id)
    except HTTPException:
        return RedirectResponse("/login", status_code=303)
    except Exception as e:
        print(f"Error loading bookings: {e}")
        bookings = []
        user = None

    return templates.TemplateResponse("my_bookings.html", {
        "request": request,
        "user": user,
        "bookings": bookings,
        "now": naive_now
    })


# --- Auth ---
@app.get("/login", response_class=HTMLResponse)
def login_get(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})


@app.post("/login")
def login_post(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    user = db.query(models.User).filter_by(username=username, is_active=True).first()
    if not user or not verify_password(password, user.hashed_password):
        return templates.TemplateResponse("login.html", {"request": request, "error": "Неверный логин или пароль"})

    # Обновляем last_login
    user.last_login = now_with_tz()
    db.commit()

    request.session["user_id"] = user.id
    return RedirectResponse("/", status_code=303)


@app.get("/register", response_class=HTMLResponse)
def register_get(request: Request):
    return templates.TemplateResponse("register.html", {"request": request})


@app.post("/register")
def register_post(request: Request, username: str = Form(...), email: str = Form(...),
                  password: str = Form(...),
                  db: Session = Depends(get_db)):
    # Проверяем существование пользователя
    if db.query(models.User).filter_by(username=username).first():
        return templates.TemplateResponse("register.html", {"request": request, "error": "Имя занято"})

    if db.query(models.User).filter_by(email=email).first():
        return templates.TemplateResponse("register.html", {"request": request, "error": "Email уже используется"})

    # Проверка email на валидность
    import re
    email_pattern = r'^[^\s@]+@[^\s@]+\.[^\s@]+$'
    if not re.match(email_pattern, email):
        return templates.TemplateResponse("register.html", {"request": request, "error": "Введите корректный email"})

    # Создаем нового пользователя
    u = models.User(
        username=username,
        email=email,
        hashed_password=hash_password(password),
        is_admin=False,
        is_active=True
    )
    db.add(u)
    db.commit()
    db.refresh(u)

    request.session["user_id"] = u.id
    return RedirectResponse("/", status_code=303)


@app.get("/logout")
def logout(request: Request):
    request.session.clear()
    return RedirectResponse("/", status_code=303)


# --- Admin ---
@app.get("/admin", response_class=HTMLResponse)
def admin_dashboard(request: Request, db: Session = Depends(get_db), user=Depends(require_admin)):
    try:
        quests = crud.get_quests(db, skip=0, limit=1000, filters={})
        all_bookings = crud.get_all_bookings(db)

        # Получаем количество активных пользователей
        active_users_count = db.query(models.User).filter(models.User.is_active == True).count()

    except Exception as e:
        quests = []
        all_bookings = []
        active_users_count = 0
        print(f"Admin error: {e}")

    return templates.TemplateResponse("admin_dashboard.html", {
        "request": request,
        "quests": quests,
        "user": user,
        "now": naive_now,
        "quest_bookings": all_bookings,
        "active_users_count": active_users_count  # Добавляем эту переменную
    })


@app.get("/admin/add", response_class=HTMLResponse)
def admin_add_form(request: Request, user=Depends(require_admin)):
    return templates.TemplateResponse("add_quest.html", {"request": request, "user": user})


@app.post("/admin/add")
async def add_post(
        request: Request,
        title: str = Form(...),
        description: str = Form(...),
        genres: List[str] = Form(...),
        difficulty: str = Form(...),
        fear_level: int = Form(...),
        max_players: int = Form(...),
        address: str = Form(...),
        price: int = Form(2000),
        image: Optional[UploadFile] = File(None),
        clipboard_image: str = Form(None),
        db: Session = Depends(get_db),
        user=Depends(require_admin)
):
    try:
        image_path = None
        image_data = None

        # Обработка изображения
        if clipboard_image and clipboard_image.startswith('data:image'):
            # Сохраняем base64 изображение в базу данных
            image_data = clipboard_image

            # Сохраняем файл локально для совместимости
            import base64
            image_bytes = base64.b64decode(clipboard_image.split(',')[1])
            safe_name = f"{uuid.uuid4().hex}.png"
            dest = os.path.join("static", "uploads", safe_name)
            with open(dest, "wb") as f:
                f.write(image_bytes)
            image_path = f"uploads/{safe_name}"

        elif image and image.filename:
            # Читаем содержимое файла и конвертируем в base64
            content = await image.read()
            import base64
            # Определяем формат изображения
            ext = os.path.splitext(image.filename)[1].lower()
            mime_type = "image/jpeg" if ext in ['.jpg', '.jpeg'] else "image/png"
            if ext == '.gif':
                mime_type = "image/gif"

            # Создаем base64 строку
            image_data = f"data:{mime_type};base64,{base64.b64encode(content).decode('utf-8')}"

            # Сохраняем файл локально
            safe_name = f"{uuid.uuid4().hex}{ext}"
            dest = os.path.join("static", "uploads", safe_name)
            with open(dest, "wb") as f:
                f.write(content)
            image_path = f"uploads/{safe_name}"

        # Объединяем выбранные жанры в строку
        genre_str = ", ".join(genres) if genres else "Не указан"

        # Создание квеста
        new_quest = models.Quest(
            title=title,
            description=description,
            genre=genre_str,
            difficulty=difficulty,
            fear_level=fear_level,
            max_players=max_players,
            min_players=2,
            address=address,
            price=price,
            image_path=image_path,
            image_data=image_data,  # Сохраняем base64 в БД
            is_active=True
        )
        db.add(new_quest)
        db.commit()

        return RedirectResponse("/admin", status_code=303)

    except Exception as e:
        print(f"Error adding quest: {e}")
        quests = crud.get_quests(db, skip=0, limit=1000, filters={})
        return templates.TemplateResponse("admin_dashboard.html", {
            "request": request,
            "quests": quests,
            "user": user,
            "error": f"Ошибка при добавлении квеста: {str(e)}",
            "now": naive_now
        })


@app.get("/admin/edit/{quest_id}", response_class=HTMLResponse)
def admin_edit_form(request: Request, quest_id: int, db: Session = Depends(get_db), user=Depends(require_admin)):
    quest = crud.get_quest(db, quest_id)
    if not quest:
        raise HTTPException(status_code=404, detail="Quest not found")
    return templates.TemplateResponse("edit_quest.html", {"request": request, "quest": quest, "user": user})


@app.post("/admin/edit/{quest_id}")
async def edit_post(
        request: Request,
        quest_id: int,
        title: str = Form(...),
        description: str = Form(...),
        genres: List[str] = Form(...),
        difficulty: str = Form(...),
        fear_level: int = Form(...),
        max_players: int = Form(...),
        address: str = Form(...),
        price: int = Form(2000),
        image: Optional[UploadFile] = File(None),
        clipboard_image: str = Form(None),
        db: Session = Depends(get_db),
        user=Depends(require_admin)
):
    quest = crud.get_quest(db, quest_id)
    if not quest:
        raise HTTPException(status_code=404, detail="Quest not found")

    # Обработка изображения
    if clipboard_image and clipboard_image.startswith('data:image'):
        # Сохраняем base64 в БД
        quest.image_data = clipboard_image

        # Сохраняем файл локально
        import base64
        image_bytes = base64.b64decode(clipboard_image.split(',')[1])
        safe_name = f"{uuid.uuid4().hex}.png"
        dest = os.path.join("static", "uploads", safe_name)
        with open(dest, "wb") as f:
            f.write(image_bytes)
        quest.image_path = f"uploads/{safe_name}"

    elif image and image.filename:
        # Читаем содержимое файла и конвертируем в base64
        content = await image.read()
        import base64
        ext = os.path.splitext(image.filename)[1].lower()
        mime_type = "image/jpeg" if ext in ['.jpg', '.jpeg'] else "image/png"
        if ext == '.gif':
            mime_type = "image/gif"

        # Сохраняем base64 в БД
        quest.image_data = f"data:{mime_type};base64,{base64.b64encode(content).decode('utf-8')}"

        # Сохраняем файл локально
        safe_name = f"{uuid.uuid4().hex}{ext}"
        dest = os.path.join("static", "uploads", safe_name)
        with open(dest, "wb") as f:
            f.write(content)
        quest.image_path = f"uploads/{safe_name}"

    # Обновляем поля
    quest.title = title
    quest.description = description
    quest.genre = ", ".join(genres) if genres else "Не указан"
    quest.difficulty = difficulty
    quest.fear_level = fear_level
    quest.max_players = max_players
    quest.address = address
    quest.price = price

    db.commit()

    return RedirectResponse("/admin", status_code=303)


@app.get("/admin/bookings", response_class=HTMLResponse)
def admin_bookings(request: Request, quest_id: Optional[int] = None, db: Session = Depends(get_db),
                   user=Depends(require_admin)):
    try:
        if quest_id:
            bookings = crud.get_quest_bookings(db, quest_id)
        else:
            bookings = crud.get_all_bookings(db)
        quests = crud.get_quests(db, skip=0, limit=1000, filters={})
    except Exception as e:
        bookings = []
        quests = []
        print(f"Bookings error: {e}")

    return templates.TemplateResponse("admin_bookings.html", {
        "request": request,
        "bookings": bookings,
        "quests": quests,
        "user": user,
        "now": naive_now
    })


@app.post("/admin/delete-booking/{booking_id}")
def admin_delete_booking(booking_id: int, db: Session = Depends(get_db), user=Depends(require_admin)):
    crud.delete_booking(db, booking_id)
    return RedirectResponse("/admin/bookings", status_code=303)


@app.post("/admin/delete/{quest_id}")
def admin_delete(request: Request, quest_id: int, db: Session = Depends(get_db), user=Depends(require_admin)):
    # Проверяем наличие будущих бронирований
    if crud.has_quest_bookings(db, quest_id):
        # Если есть бронирования, показываем страницу с предупреждением
        quest = crud.get_quest(db, quest_id)
        bookings = crud.get_quest_bookings(db, quest_id)
        quests = crud.get_quests(db, skip=0, limit=1000, filters={})
        return templates.TemplateResponse("admin_dashboard.html", {
            "request": request,
            "quests": quests,
            "user": user,
            "error": f"Невозможно удалить квест '{quest.title}'. У него есть активные бронирования.",
            "blocked_quest_id": quest_id,
            "quest_bookings": bookings,
            "now": naive_now
        })

    # ПОЛНОЕ УДАЛЕНИЕ из БД
    crud.delete_quest(db, quest_id)
    return RedirectResponse("/admin", status_code=303)


@app.post("/admin/delete-quest-with-bookings/{quest_id}")
def admin_delete_quest_with_bookings(quest_id: int, db: Session = Depends(get_db), user=Depends(require_admin)):
    # Принудительное удаление квеста вместе со всеми бронированиями
    crud.delete_quest(db, quest_id)
    return RedirectResponse("/admin", status_code=303)


@app.post("/admin/delete-all-bookings/{quest_id}")
def admin_delete_all_bookings(quest_id: int, db: Session = Depends(get_db), user=Depends(require_admin)):
    bookings = crud.get_quest_bookings(db, quest_id)
    for booking in bookings:
        crud.delete_booking(db, booking.id)
    return RedirectResponse(f"/admin?highlight={quest_id}", status_code=303)


# Страница с отзывами на квест
@app.get("/quest/{quest_id}/reviews", response_class=HTMLResponse)
def quest_reviews(request: Request, quest_id: int, db: Session = Depends(get_db)):
    quest = crud.get_quest(db, quest_id)
    if not quest:
        raise HTTPException(status_code=404, detail="Quest not found")

    reviews = get_quest_reviews(db, quest_id)
    avg_rating = get_quest_average_rating(db, quest_id)

    try:
        user = get_current_user(request, db)
        can_review = can_user_review_quest(db, user.id, quest_id)
    except:
        user = None
        can_review = False

    return templates.TemplateResponse("quest_reviews.html", {
        "request": request,
        "quest": quest,
        "reviews": reviews,
        "avg_rating": avg_rating,
        "user": user,
        "can_review": can_review,
        "now": naive_now
    })


# Страница для написания отзыва
@app.get("/write-review/{booking_id}", response_class=HTMLResponse)
def write_review_form(request: Request, booking_id: int, db: Session = Depends(get_db)):
    try:
        user = get_current_user(request, db)

        booking = db.query(models.Booking).filter(
            models.Booking.id == booking_id,
            models.Booking.user_id == user.id
        ).first()

        if not booking:
            return RedirectResponse("/my-bookings", status_code=303)

        # Проверяем, что квест уже прошел
        if booking.booking_date_time.replace(tzinfo=None) >= naive_now():
            return RedirectResponse("/my-bookings", status_code=303)

        # Проверяем, что предоплата внесена
        if booking.payment_status != 'prepayment_paid':
            return RedirectResponse("/my-bookings", status_code=303)

        # Проверяем, нет ли уже отзыва
        existing = db.query(models.Review).filter(
            models.Review.booking_id == booking_id
        ).first()

        if existing:
            return RedirectResponse(f"/quest/{booking.quest_id}/reviews", status_code=303)

        return templates.TemplateResponse("write_review.html", {
            "request": request,
            "booking": booking,
            "quest": booking.quest,
            "user": user
        })

    except HTTPException:
        return RedirectResponse("/login", status_code=303)
    except Exception as e:
        print(f"Write review error: {e}")
        import traceback
        traceback.print_exc()
        return RedirectResponse("/my-bookings", status_code=303)


# Обработка отправки отзыва
@app.post("/submit-review")
async def submit_review(request: Request, booking_id: int = Form(...),
                        rating: int = Form(...), comment: str = Form(""),
                        db: Session = Depends(get_db)):
    try:
        user = get_current_user(request, db)

        booking = db.query(models.Booking).filter(
            models.Booking.id == booking_id,
            models.Booking.user_id == user.id
        ).first()

        if not booking:
            return JSONResponse({"success": False, "message": "Бронирование не найдено"}, status_code=404)

        # Проверяем, что квест уже прошел
        if booking.booking_date_time.replace(tzinfo=None) >= naive_now():
            return JSONResponse({"success": False, "message": "Отзыв можно оставить только после прохождения квеста"}, status_code=400)

        # Проверяем, что предоплата внесена
        if booking.payment_status != 'prepayment_paid':
            return JSONResponse({"success": False, "message": "Отзыв можно оставить только после оплаты"}, status_code=400)

        # Проверяем, не оставлял ли пользователь уже отзыв на это бронирование
        existing_review = db.query(models.Review).filter(
            models.Review.booking_id == booking_id
        ).first()

        if existing_review:
            return JSONResponse({"success": False, "message": "Вы уже оставили отзыв на это бронирование"}, status_code=400)

        review = create_review(db, user.id, booking.quest_id, booking_id, rating, comment)

        if not review:
            return JSONResponse({"success": False, "message": "Нельзя оставить отзыв"}, status_code=400)

        return JSONResponse({"success": True, "message": "Отзыв успешно добавлен!"})

    except HTTPException:
        return JSONResponse({"success": False, "message": "Необходимо авторизоваться"}, status_code=401)
    except Exception as e:
        print(f"Review error: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse({"success": False, "message": f"Ошибка: {str(e)}"}, status_code=500)


# --- Документы ---
@app.post("/download-receipt")
async def download_receipt(request: Request, db: Session = Depends(get_db)):
    try:
        user = get_current_user(request, db)
    except:
        user = None

    data = await request.json()
    buffer = io.BytesIO()

    receipt_text = f"""ЧЕК О БРОНИРОВАНИИ
══════════════════════════════════════════════════
Квест: {data['quest_title']}
──────────────────────────────────────────────────
Полная стоимость: {data['quest_price']}₽
Предоплата (50%): {data['prepayment']}₽
Способ оплаты: {data['payment_method']}
Статус оплаты: {data['payment_status']}
──────────────────────────────────────────────────
Клиент: {user.username if user else 'Гость'}
Email: {user.email if user else 'Не указан'}
Дата бронирования: {naive_now().strftime('%d.%m.%Y %H:%M')}
──────────────────────────────────────────────────
Условия:
✓ Оставшуюся сумму оплатите на месте
✓ При отмене менее чем за 24 часа
  предоплата не возвращается
══════════════════════════════════════════════════
Спасибо за бронирование!"""

    buffer.write(receipt_text.encode('utf-8'))
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename=receipt_{naive_now().strftime('%Y%m%d_%H%M%S')}.txt"}
    )


@app.post("/cancel-booking/{booking_id}")
def cancel_booking(request: Request, booking_id: int, db: Session = Depends(get_db)):
    try:
        user = get_current_user(request, db)

        booking = db.query(models.Booking).filter(
            models.Booking.id == booking_id,
            models.Booking.user_id == user.id
        ).first()

        if not booking:
            return JSONResponse({"success": False, "message": "Бронирование не найдено"}, status_code=404)

        booking_time = booking.booking_date_time.replace(tzinfo=None)
        current_time = naive_now()

        if booking_time < current_time:
            return JSONResponse({"success": False, "message": "Нельзя отменить прошедшее бронирование"},
                                status_code=400)

        time_diff = (booking_time - current_time).total_seconds() / 3600
        if time_diff < 24:
            return JSONResponse(
                {"success": False, "message": "Отмена возможна не менее чем за 24 часа до начала квеста"},
                status_code=400)

        quest_title = booking.quest.title
        booking_date = booking.booking_date_time.strftime('%d.%m.%Y')
        booking_time_str = booking.booking_date_time.strftime('%H:%M')

        crud.delete_booking(db, booking_id)

        try:
            send_booking_cancellation(
                user_email=user.email,
                user_name=user.username,
                quest_title=quest_title,
                booking_date=booking_date,
                booking_time=booking_time_str
            )
        except Exception as e:
            print(f"Email sending error: {e}")

        return JSONResponse(
            {"success": True, "message": "Бронирование успешно отменено. Уведомление отправлено на email."})

    except HTTPException:
        return JSONResponse({"success": False, "message": "Необходимо авторизоваться"}, status_code=401)
    except Exception as e:
        print(f"Cancel booking error: {e}")
        return JSONResponse({"success": False, "message": f"Ошибка: {str(e)}"}, status_code=500)


@app.post("/pay-prepayment/{booking_id}")
def pay_prepayment(request: Request, booking_id: int, db: Session = Depends(get_db)):
    try:
        user = get_current_user(request, db)

        booking = db.query(models.Booking).filter(
            models.Booking.id == booking_id,
            models.Booking.user_id == user.id
        ).first()

        if not booking:
            return JSONResponse({"success": False, "message": "Бронирование не найдено"}, status_code=404)

        if booking.payment_status == 'prepayment_paid':
            return JSONResponse({"success": False, "message": "Предоплата уже внесена"}, status_code=400)

        booking_time = booking.booking_date_time.replace(tzinfo=None)
        current_time = naive_now()

        if booking_time < current_time:
            return JSONResponse({"success": False, "message": "Нельзя оплатить прошедшее бронирование"},
                                status_code=400)

        booking.payment_status = 'prepayment_paid'
        db.commit()

        quest = booking.quest
        booking_datetime = booking.booking_date_time
        booking_date = booking_datetime.strftime('%d.%m.%Y')
        booking_time_str = booking_datetime.strftime('%H:%M')

        try:
            send_booking_confirmation(
                user_email=user.email,
                user_name=user.username,
                quest_title=quest.title,
                booking_date=booking_date,
                booking_time=booking_time_str,
                address=quest.address,
                prepayment=booking.prepayment,
                total_price=booking.total_price
            )
        except Exception as e:
            print(f"Email sending error: {e}")

        return JSONResponse({
            "success": True,
            "message": f"Предоплата {booking.prepayment}₽ успешно внесена! На ваш email отправлено подтверждение.",
            "prepayment": booking.prepayment,
            "total_price": booking.total_price
        })

    except HTTPException:
        return JSONResponse({"success": False, "message": "Необходимо авторизоваться"}, status_code=401)
    except Exception as e:
        print(f"Pay prepayment error: {e}")
        return JSONResponse({"success": False, "message": f"Ошибка: {str(e)}"}, status_code=500)


@app.get("/admin/report/word")
async def report_word(db: Session = Depends(get_db), user=Depends(require_admin)):
    """Генерация отчета по бронированиям в Word с логотипом и печатью"""
    try:
        bookings = crud.get_all_bookings(db)

        doc = Document()

        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

        # Создаем таблицу для шапки с логотипом
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(1.5)
        header_table.columns[1].width = Inches(4.5)

        # Добавляем логотип в первую ячейку
        try:
            logo_path = "static/images/logo_black.png"
            if os.path.exists(logo_path):
                logo_cell = header_table.cell(0, 0)
                logo_paragraph = logo_cell.paragraphs[0]
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(logo_path, width=Inches(1.8), height=Inches(1.8))
        except:
            pass

        # Добавляем информацию во вторую ячейку
        info_cell = header_table.cell(0, 1)
        info_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_cell.paragraphs[0].add_run("Алиби\n").bold = True
        info_cell.paragraphs[0].add_run("г.Москва, ул.Квестовая, д.88\n")
        info_cell.paragraphs[0].add_run("e-mail: alibi.quest@mail.ru")

        doc.add_paragraph()

        # Заголовок отчета
        report_title = doc.add_paragraph("ОТЧЕТ ПО БРОНИРОВАНИЯМ")
        report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_title.runs[0].bold = True
        report_title.runs[0].font.size = Pt(14)

        # Дата формирования
        date_para = doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Создаем таблицу с данными
        if bookings:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            table.autofit = False

            # Устанавливаем ширину колонок
            col_widths = [0.5, 1.5, 2.0, 1.2, 1.2, 1.0]
            for i, width in enumerate(col_widths):
                table.columns[i].width = Inches(width)

            # Заголовки таблицы
            headers = ['№', 'Пользователь', 'Квест', 'Дата и время', 'Сумма', 'Статус']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Заливаем фон заголовков
                shading_elm = parse_xml(r'<w:shd {} w:fill="E6E6FA"/>'.format(nsdecls('w')))
                hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

            # Данные
            total_revenue = 0
            for i, booking in enumerate(bookings, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_cells[1].text = booking.user.username if booking.user else 'Не указан'
                row_cells[2].text = booking.quest.title

                booking_time = booking.booking_date_time.strftime(
                    '%d.%m.%Y %H:%M') if booking.booking_date_time else 'Не указано'
                row_cells[3].text = booking_time
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_cells[4].text = f"{booking.total_price}₽"
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                status_text = "Оплачен" if booking.payment_status == 'prepayment_paid' else "Ожидает оплату"
                row_cells[5].text = status_text
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                total_revenue += booking.total_price

            doc.add_paragraph()

            # Итоги
            total_para = doc.add_paragraph()
            total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total_para.add_run("═" * 50 + "\n").bold = True
            total_para.add_run(f"ИТОГО БРОНИРОВАНИЙ: {len(bookings)}\n").bold = True
            total_para.add_run(f"ОБЩАЯ ВЫРУЧКА: {total_revenue} руб\n").bold = True
            total_para.add_run("═" * 50).bold = True

        else:
            doc.add_paragraph("Нет данных о бронированиях")

        doc.add_paragraph()
        doc.add_paragraph()

        # Создаем таблицу для подписи и печати
        footer_table = doc.add_table(rows=1, cols=2)
        footer_table.autofit = False
        footer_table.columns[0].width = Inches(4.0)
        footer_table.columns[1].width = Inches(2.0)

        # Подпись в левой ячейке
        sign_cell = footer_table.cell(0, 0)
        sign_cell.paragraphs[0].add_run("_________________________\n")
        sign_cell.paragraphs[0].add_run("Подпись ответственного лица")

        # Печать в правой ячейке
        try:
            stamp_path = "static/images/stamp.png"
            if os.path.exists(stamp_path):
                stamp_cell = footer_table.cell(0, 1)
                stamp_paragraph = stamp_cell.paragraphs[0]
                stamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                stamp_run = stamp_paragraph.add_run()
                stamp_run.add_picture(stamp_path, width=Inches(1.8), height=Inches(1.8))
        except:
            pass

        # Сохраняем в буфер
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"otchet_bronirovaniya_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        encoded_filename = urllib.parse.quote(filename)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )

    except ImportError:
        return JSONResponse({"message": "Для генерации Word отчетов установите python-docx: pip install python-docx"})
    except Exception as e:
        print(f"Error generating bookings report: {e}")
        return JSONResponse({"message": f"Ошибка при генерации отчета: {str(e)}"}, status_code=500)


@app.get("/admin/report/quests")
async def report_quests_word(db: Session = Depends(get_db), user=Depends(require_admin)):
    """Генерация отчета по квестам в Word с рейтингом и критериями оценки"""
    try:
        # Получаем все активные квесты
        quests = crud.get_quests(db, skip=0, limit=1000, filters={})

        # Получаем статистику по каждому квесту
        quest_stats = []
        total_quests_price = 0

        for quest in quests:
            # Количество бронирований для этого квеста
            bookings_count = db.query(models.Booking).filter(
                models.Booking.quest_id == quest.id
            ).count()

            # Количество завершенных бронирований (прошедших)
            completed_bookings = db.query(models.Booking).filter(
                models.Booking.quest_id == quest.id,
                models.Booking.booking_date_time < datetime.now()
            ).count()

            # Доход от этого квеста
            revenue = db.query(func.sum(models.Booking.total_price)).filter(
                models.Booking.quest_id == quest.id
            ).scalar() or 0

            # Средний рейтинг квеста
            avg_rating = get_quest_average_rating(db, quest.id)

            # Количество отзывов
            reviews_count = db.query(models.Review).filter(
                models.Review.quest_id == quest.id
            ).count()

            # Процент заполняемости (если есть бронирования)
            fill_rate = 0
            if bookings_count > 0:
                # Предполагаем максимальную вместимость в день (2 слота в день * max_players)
                max_capacity = quest.max_players * 2  # примерно
                fill_rate = min(100, int((bookings_count / max_capacity) * 100)) if max_capacity > 0 else 0

            quest_stats.append({
                'quest': quest,
                'bookings_count': bookings_count,
                'completed_bookings': completed_bookings,
                'revenue': revenue,
                'avg_rating': avg_rating,
                'reviews_count': reviews_count,
                'fill_rate': fill_rate
            })

            total_quests_price += quest.price

        doc = Document()

        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

        # Создаем таблицу для шапки с логотипом
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(1.5)
        header_table.columns[1].width = Inches(4.5)

        # Добавляем логотип в первую ячейку
        try:
            logo_path = "static/images/logo_black.png"
            if os.path.exists(logo_path):
                logo_cell = header_table.cell(0, 0)
                logo_paragraph = logo_cell.paragraphs[0]
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(logo_path, width=Inches(1.8), height=Inches(1.8))
        except:
            pass

        # Добавляем информацию во вторую ячейку
        info_cell = header_table.cell(0, 1)
        info_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_cell.paragraphs[0].add_run("Алиби\n").bold = True
        info_cell.paragraphs[0].add_run("г.Москва, ул.Квестовая, д.88\n")
        info_cell.paragraphs[0].add_run("e-mail: alibi.quest@mail.ru")

        doc.add_paragraph()

        # Заголовок отчета
        report_title = doc.add_paragraph("ОТЧЕТ ПО КВЕСТАМ С РЕЙТИНГОМ И КРИТЕРИЯМИ ОЦЕНКИ")
        report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_title.runs[0].bold = True
        report_title.runs[0].font.size = Pt(14)

        # Дата формирования
        date_para = doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Создаем таблицу с данными о квестах
        if quests:
            # Таблица с основными данными (10 колонок)
            table = doc.add_table(rows=1, cols=10)
            table.style = 'Table Grid'
            table.autofit = False

            # Устанавливаем ширину колонок
            col_widths = [0.5, 1.8, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8]
            for i, width in enumerate(col_widths):
                table.columns[i].width = Inches(width)

            # Заголовки таблицы
            headers = ['№', 'Название', 'Жанр', 'Сложность', 'Страх', 'Цена', 'Броней', 'Доход', 'Рейтинг', 'Отзывов']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Заливаем фон заголовков
                shading_elm = parse_xml(r'<w:shd {} w:fill="E6E6FA"/>'.format(nsdecls('w')))
                hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

            # Данные
            total_revenue = 0
            total_bookings = 0
            total_ratings_sum = 0
            quests_with_rating = 0

            for i, stats in enumerate(quest_stats, 1):
                quest = stats['quest']
                row_cells = table.add_row().cells

                row_cells[0].text = str(i)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_cells[1].text = quest.title

                row_cells[2].text = quest.genre
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_cells[3].text = quest.difficulty
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_cells[4].text = f"{quest.fear_level}/5"
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_cells[5].text = f"{quest.price}₽"
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_cells[6].text = str(stats['bookings_count'])
                row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_cells[7].text = f"{stats['revenue']}₽"
                row_cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Рейтинг со звездочками
                if stats['avg_rating']:
                    rating_text = f"{stats['avg_rating']} ★"
                    if stats['avg_rating'] >= 4.5:
                        rating_text += " (Отлично)"
                    elif stats['avg_rating'] >= 3.5:
                        rating_text += " (Хорошо)"
                    elif stats['avg_rating'] >= 2.5:
                        rating_text += " (Средне)"
                    else:
                        rating_text += " (Низкий)"
                    row_cells[8].text = rating_text
                    total_ratings_sum += stats['avg_rating']
                    quests_with_rating += 1
                else:
                    row_cells[8].text = "Нет оценок"

                row_cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_cells[9].text = str(stats['reviews_count'])
                row_cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                total_revenue += stats['revenue']
                total_bookings += stats['bookings_count']

            doc.add_paragraph()

            # Общая статистика
            avg_rating_all = round(total_ratings_sum / quests_with_rating, 1) if quests_with_rating > 0 else 0

            summary_para = doc.add_paragraph()
            summary_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            summary_para.add_run("═" * 60 + "\n").bold = True
            summary_para.add_run(f"ИТОГО КВЕСТОВ: {len(quests)}\n").bold = True
            summary_para.add_run(f"ИТОГО БРОНИРОВАНИЙ: {total_bookings}\n").bold = True
            summary_para.add_run(f"ИТОГОВАЯ ВЫРУЧКА: {total_revenue} руб\n").bold = True
            summary_para.add_run(
                f"СРЕДНЯЯ ЦЕНА КВЕСТА: {total_quests_price // len(quests) if quests else 0} руб\n").bold = True
            summary_para.add_run(f"СРЕДНИЙ РЕЙТИНГ ВСЕХ КВЕСТОВ: {avg_rating_all} ★\n").bold = True
            summary_para.add_run("═" * 60).bold = True

            # Дополнительная статистика по жанрам
            doc.add_paragraph()
            genre_stats_title = doc.add_paragraph("Статистика по жанрам:")
            genre_stats_title.runs[0].bold = True

            genres_count = {}
            genres_rating = {}
            for stats in quest_stats:
                quest = stats['quest']
                genre = quest.genre
                for g in genre.split(', '):
                    genres_count[g] = genres_count.get(g, 0) + 1
                    if stats['avg_rating']:
                        if g not in genres_rating:
                            genres_rating[g] = {'sum': 0, 'count': 0}
                        genres_rating[g]['sum'] += stats['avg_rating']
                        genres_rating[g]['count'] += 1

            genre_table = doc.add_table(rows=1, cols=3)
            genre_table.style = 'Table Grid'
            genre_table.columns[0].width = Inches(3.0)
            genre_table.columns[1].width = Inches(1.5)
            genre_table.columns[2].width = Inches(2.0)

            genre_header = genre_table.rows[0].cells
            genre_header[0].text = "Жанр"
            genre_header[1].text = "Количество"
            genre_header[2].text = "Средний рейтинг"
            genre_header[0].paragraphs[0].runs[0].bold = True
            genre_header[1].paragraphs[0].runs[0].bold = True
            genre_header[2].paragraphs[0].runs[0].bold = True

            for genre, count in sorted(genres_count.items(), key=lambda x: x[1], reverse=True):
                row = genre_table.add_row().cells
                row[0].text = genre
                row[1].text = str(count)
                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                if genre in genres_rating:
                    avg = round(genres_rating[genre]['sum'] / genres_rating[genre]['count'], 1)
                    row[2].text = f"{avg} ★"
                else:
                    row[2].text = "Нет оценок"
                row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            # Таблица рейтинга квестов
            rating_title = doc.add_paragraph("РЕЙТИНГ КВЕСТОВ (ТОП-5):")
            rating_title.runs[0].bold = True
            rating_title.runs[0].font.size = Pt(12)

            # Сортируем по рейтингу
            top_quests = sorted([s for s in quest_stats if s['avg_rating']],
                                key=lambda x: x['avg_rating'], reverse=True)[:5]

            if top_quests:
                rating_table = doc.add_table(rows=1, cols=4)
                rating_table.style = 'Table Grid'
                rating_table.columns[0].width = Inches(0.8)
                rating_table.columns[1].width = Inches(3.0)
                rating_table.columns[2].width = Inches(1.2)
                rating_table.columns[3].width = Inches(1.5)

                rating_headers = ['Место', 'Название квеста', 'Рейтинг', 'Кол-во отзывов']
                rating_hdr = rating_table.rows[0].cells
                for i, header in enumerate(rating_headers):
                    rating_hdr[i].text = header
                    rating_hdr[i].paragraphs[0].runs[0].bold = True
                    rating_hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for idx, stats in enumerate(top_quests, 1):
                    row = rating_table.add_row().cells
                    row[0].text = str(idx)
                    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row[1].text = stats['quest'].title
                    row[2].text = f"{stats['avg_rating']} ★"
                    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row[3].text = str(stats['reviews_count'])
                    row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph("Нет данных для формирования рейтинга")

            doc.add_paragraph()

            # Таблица популярности квестов (по количеству бронирований)
            popular_title = doc.add_paragraph("ПОПУЛЯРНОСТЬ КВЕСТОВ (ТОП-5):")
            popular_title.runs[0].bold = True
            popular_title.runs[0].font.size = Pt(12)

            top_popular = sorted(quest_stats, key=lambda x: x['bookings_count'], reverse=True)[:5]

            if top_popular and any(s['bookings_count'] > 0 for s in top_popular):
                popular_table = doc.add_table(rows=1, cols=4)
                popular_table.style = 'Table Grid'
                popular_table.columns[0].width = Inches(0.8)
                popular_table.columns[1].width = Inches(3.0)
                popular_table.columns[2].width = Inches(1.2)
                popular_table.columns[3].width = Inches(1.5)

                popular_headers = ['Место', 'Название квеста', 'Бронирований', 'Доход']
                popular_hdr = popular_table.rows[0].cells
                for i, header in enumerate(popular_headers):
                    popular_hdr[i].text = header
                    popular_hdr[i].paragraphs[0].runs[0].bold = True
                    popular_hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for idx, stats in enumerate(top_popular, 1):
                    if stats['bookings_count'] == 0:
                        continue
                    row = popular_table.add_row().cells
                    row[0].text = str(idx)
                    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row[1].text = stats['quest'].title
                    row[2].text = str(stats['bookings_count'])
                    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row[3].text = f"{stats['revenue']}₽"
                    row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph("Нет данных о бронированиях")

        else:
            doc.add_paragraph("Нет данных о квестах")

        doc.add_paragraph()
        doc.add_paragraph()

        # Создаем таблицу для подписи и печати
        footer_table = doc.add_table(rows=1, cols=2)
        footer_table.autofit = False
        footer_table.columns[0].width = Inches(4.0)
        footer_table.columns[1].width = Inches(2.0)

        # Подпись в левой ячейке
        sign_cell = footer_table.cell(0, 0)
        sign_cell.paragraphs[0].add_run("_________________________\n")
        sign_cell.paragraphs[0].add_run("Подпись ответственного лица")

        # Печать в правой ячейке
        try:
            stamp_path = "static/images/stamp.png"
            if os.path.exists(stamp_path):
                stamp_cell = footer_table.cell(0, 1)
                stamp_paragraph = stamp_cell.paragraphs[0]
                stamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                stamp_run = stamp_paragraph.add_run()
                stamp_run.add_picture(stamp_path, width=Inches(1.8), height=Inches(1.8))
        except:
            pass

        # Сохраняем в буфер
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"otchet_kvesty_s_reitingom_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        encoded_filename = urllib.parse.quote(filename)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )

    except ImportError:
        return JSONResponse({"message": "Для генерации Word отчетов установите python-docx: pip install python-docx"})
    except Exception as e:
        print(f"Error generating quests report: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse({"message": f"Ошибка при генерации отчета: {str(e)}"}, status_code=500)


@app.get("/api/quest-image/{quest_id}")
async def get_quest_image(quest_id: int, db: Session = Depends(get_db)):
    """Возвращает изображение квеста из базы данных"""
    import base64
    from fastapi.responses import Response

    quest = crud.get_quest(db, quest_id)
    if not quest:
        raise HTTPException(status_code=404, detail="Quest not found")

    # Проверяем наличие изображения в БД
    if quest.image_data:
        try:
            # Формат: data:image/png;base64,xxxxx
            header, encoded = quest.image_data.split(',', 1)
            mime_type = header.split(':')[1].split(';')[0]
            image_bytes = base64.b64decode(encoded)
            return Response(content=image_bytes, media_type=mime_type)
        except Exception as e:
            print(f"Error serving image from DB: {e}")

    # Если нет в БД, пробуем загрузить из файла
    if quest.image_path:
        file_path = os.path.join("static", quest.image_path)
        if os.path.exists(file_path):
            with open(file_path, "rb") as f:
                content = f.read()
            ext = os.path.splitext(file_path)[1].lower()
            mime_type = "image/jpeg" if ext in ['.jpg', '.jpeg'] else "image/png"
            if ext == '.gif':
                mime_type = "image/gif"
            return Response(content=content, media_type=mime_type)

    raise HTTPException(status_code=404, detail="Image not found")

# --- API для проверки ---
@app.get("/api/quest-has-bookings/{quest_id}")
def api_quest_has_bookings(quest_id: int, db: Session = Depends(get_db)):
    has_bookings = crud.has_quest_bookings(db, quest_id)
    return JSONResponse({"has_bookings": has_bookings})


@app.get("/api/quests")
def api_get_quests(
        request: Request,  # <-- Первый параметр
        skip: int = 0,
        q: Optional[str] = None,
        genre: Optional[List[str]] = Query(None),  # <-- Важно: Query(None)
        difficulty: Optional[List[str]] = Query(None),  # <-- Важно: Query(None)
        fear_level: Optional[str] = None,
        players: Optional[str] = None,
        sort: Optional[str] = None,
        db: Session = Depends(get_db)
):
    print(f"API DEBUG - genre: {genre}")
    print(f"API DEBUG - difficulty: {difficulty}")

    filters = {
        "q": q,
        "genre": genre,
        "difficulty": difficulty,
        "fear_level": fear_level,
        "players": players,
        "sort": sort
    }
    quests = crud.get_quests(db, skip=skip, limit=15, filters=filters)
    return templates.TemplateResponse("_quest_cards.html", {"request": request, "quests": quests})

if __name__ == "__main__":
    uvicorn.run("main:app", host="127.0.0.1", port=5000, reload=True)